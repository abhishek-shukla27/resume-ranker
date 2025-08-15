from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

def build_template_resume(data):
    """
    Builds a DOCX resume from structured resume data.
    """
    doc = Document()

    # ===== NAME =====
    if data.get("name"):
        name_para = doc.add_paragraph(data["name"])
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(16)

    # ===== CONTACT =====
    if data.get("contact"):
        contact_para = doc.add_paragraph(data["contact"])
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== SUMMARY =====
    if data.get("summary"):
        add_heading_with_line(doc, "SUMMARY")
        doc.add_paragraph(data["summary"])

    # ===== SKILLS =====
    if data.get("skills"):
        add_heading_with_line(doc, "SKILLS")
        skills_str = ", ".join([s for s in data["skills"] if s.strip()])
        if skills_str:
            doc.add_paragraph(skills_str)

    # ===== EXPERIENCE =====
    if data.get("experience"):
        experience_entries = [exp for exp in data["experience"] if exp.get("role") or exp.get("company")]
        if experience_entries:
            add_heading_with_line(doc, "EXPERIENCE")
            for exp in experience_entries:
                role_line = f"{exp.get('role', '')} – {exp.get('company', '')} ({exp.get('duration', '')})".strip()
                if role_line:
                    role_para = doc.add_paragraph(role_line)
                    role_para.runs[0].bold = True
                for d in exp.get("details", []):
                    if d.strip():
                        doc.add_paragraph(d, style="List Bullet")

    # ===== PROJECTS =====
    if data.get("projects"):
        project_entries = [proj for proj in data["projects"] if proj.get("name")]
        if project_entries:
            add_heading_with_line(doc, "PROJECTS")
            for proj in _format_projects(project_entries):
                proj_para = doc.add_paragraph(proj["name"].upper())
                proj_para.runs[0].bold = True
                for bullet in proj["details"]:
                    if bullet.strip():
                        doc.add_paragraph(bullet, style="List Bullet")

    # ===== EDUCATION =====
    if data.get("education"):
        add_heading_with_line(doc, "EDUCATION")
        if isinstance(data["education"], list):
            top_two = [e for e in data["education"] if str(e).strip()][:2]
            for e in top_two:
                doc.add_paragraph(str(e).strip(), style="List Bullet")
        else:
            edu_text = str(data["education"]).strip()
            if edu_text:
                doc.add_paragraph(edu_text, style="List Bullet")

    # ===== CERTIFICATIONS =====
    if data.get("certifications"):
        certs = [c for c in data["certifications"] if str(c).strip()]
        if certs:
            add_heading_with_line(doc, "CERTIFICATIONS")
            for cert in certs:
                doc.add_paragraph(str(cert).strip(), style="List Bullet")

    # ===== SAVE TO BYTES =====
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def add_heading_with_line(doc, text):
    """
    Adds a heading with an underline/border.
    """
    para = doc.add_paragraph()
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)

    # Add border
    p_pr = para._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_borders.append(bottom)
    p_pr.append(p_borders)


def _format_projects(projects):
    """
    Ensure projects have Objective, Tech Stack, Features.
    """
    formatted = []
    for proj in projects:
        name = proj.get("name", "Untitled Project")
        tech = proj.get("tech", "")
        details = proj.get("details", [])

        bullets = []
        if details:
            if not details[0].lower().startswith("objective"):
                bullets.append(f"Objective: {details[0]}")
            else:
                bullets.append(details[0])

        if tech:
            bullets.append(f"Tech Stack: {tech}")

        if len(details) > 1:
            if not details[1].lower().startswith("features"):
                bullets.append(f"Features: {details[1]}")
            else:
                bullets.append(details[1])

        formatted.append({"name": name, "details": bullets})
    return formatted
