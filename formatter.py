# formatter.py
from docx import Document
from io import BytesIO

def generate_docx_from_text(resume_text):
    doc = Document()
    for line in resume_text.split("\n"):
        if line.strip():
            if ":" in line and len(line) < 50:  # heading detection
                doc.add_heading(line.strip(), level=2)
            else:
                doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
